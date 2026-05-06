"""文件上传 API 路由 - docx解析 / 飞书文档读取"""
import os
import httpx
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
from typing import Optional

router = APIRouter(prefix="/api/upload", tags=["upload"])


class FeishuRequest(BaseModel):
    url: Optional[str] = None
    doc_token: Optional[str] = None
    app_id: Optional[str] = None
    app_secret: Optional[str] = None


@router.post("/docx")
async def upload_docx(file: UploadFile = File(...)):
    """上传 .docx 文件，解析并返回提取的文本内容"""
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式文件")

    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装")

    try:
        content = await file.read()
        import io
        doc = docx.Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        result = '\n\n'.join(paragraphs)
        return {
            "ok": True,
            "filename": file.filename,
            "content": result,
            "length": len(result)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"解析失败: {str(e)}")


@router.post("/feishu")
async def upload_feishu(req: FeishuRequest):
    """读取飞书文档内容

    支持两种方式：
    1. 传入完整的飞书文档 URL（如 https://feishu.cn/docx/XXX）
    2. 直接传入 doc_token
    """
    doc_token = req.doc_token
    url = req.url or ""

    # 从 URL 中提取 doc_token
    if not doc_token and url:
        # 支持多种飞书文档 URL 格式
        import re
        match = re.search(r'/docx/([a-zA-Z0-9]+)', url)
        if not match:
            match = re.search(r'/doc/([a-zA-Z0-9]+)', url)
        if not match:
            # 尝试最后一段作为 token
            parts = url.rstrip('/').split('/')
            if parts:
                doc_token = parts[-1]
        else:
            doc_token = match.group(1)

    if not doc_token:
        raise HTTPException(status_code=400, detail="请提供飞书文档 URL 或 doc_token")

    # 获取飞书 API 凭证
    app_id = req.app_id or os.getenv("FEISHU_APP_ID", "")
    app_secret = req.app_secret or os.getenv("FEISHU_APP_SECRET", "")

    if not app_id or not app_secret:
        raise HTTPException(
            status_code=500,
            detail="飞书 API 凭证未配置，请设置 FEISHU_APP_ID 和 FEISHU_APP_SECRET 环境变量"
        )

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            # 1. 获取 tenant_access_token
            token_resp = await client.post(
                "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
                json={"app_id": app_id, "app_secret": app_secret}
            )
            token_data = token_resp.json()
            if token_data.get("code") != 0:
                raise HTTPException(
                    status_code=500,
                    detail=f"获取飞书 Token 失败: {token_data.get('msg', '未知错误')}"
                )
            tenant_token = token_data["tenant_access_token"]

            # 2. 读取文档纯文本内容
            doc_resp = await client.get(
                f"https://open.feishu.cn/open-apis/docx/v1/documents/{doc_token}/raw_content",
                headers={"Authorization": f"Bearer {tenant_token}"}
            )
            doc_data = doc_resp.json()

            if doc_data.get("code") != 0:
                raise HTTPException(
                    status_code=500,
                    detail=f"读取飞书文档失败: {doc_data.get('msg', '未知错误')}"
                )

            content = doc_data.get("data", {}).get("content", "")
            return {
                "ok": True,
                "doc_token": doc_token,
                "content": content,
                "length": len(content)
            }

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="请求飞书 API 超时")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取飞书文档失败: {str(e)}")
