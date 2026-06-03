"""Export API — generate TXT, EPUB, DOCX, HTML, PDF downloads for a novel."""

import io
import logging
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.security import get_current_approved_user
from app.db.session import get_db
from app.models.chapter import Chapter
from app.models.novel import Novel
from app.models.user import User

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/novels/{novel_id}/export", tags=["export"])

VALID_FORMATS = {"txt", "epub", "docx", "pdf", "html"}
CONTENT_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "epub": "application/epub+zip",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "html": "text/html; charset=utf-8",
}
EXTENSIONS = {"txt": "txt", "epub": "epub", "docx": "docx", "pdf": "pdf", "html": "html"}


def _make_filename(title: str, fmt: str) -> str:
    ts = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    safe_title = title.strip().replace(" ", "_")[:60]
    return f"{safe_title}_{ts}.{EXTENSIONS[fmt]}"


def _disposition(filename: str) -> str:
    from urllib.parse import quote
    return f'attachment; filename*=UTF-8\'\'{quote(filename)}'


# ── TXT ───────────────────────────────────────────────

async def _generate_txt(novel: Novel, chapters: List[Chapter]) -> io.BytesIO:
    buf = io.BytesIO()
    buf.write(f"《{novel.title}》\n".encode("utf-8"))
    buf.write(f"作者：{novel.author.username}\n\n".encode("utf-8"))

    for ch in chapters:
        buf.write(f"=== {ch.title} ===\n".encode("utf-8"))
        content = (ch.content or "").strip()
        if content:
            buf.write(content.encode("utf-8"))
            buf.write(b"\n")
        buf.write(b"\n")

    buf.seek(0)
    return buf


# ── EPUB ──────────────────────────────────────────────

async def _generate_epub(novel: Novel, chapters: List[Chapter]) -> io.BytesIO:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(f"novel-{novel.id}-export")
    book.set_title(novel.title)
    book.set_language("zh")
    book.add_author(novel.author.username)

    epub_items: List[epub.EpubHtml] = []
    spine = ["nav"]

    for ch in chapters:
        ch_title = ch.title or f"Chapter {ch.order_index}"
        ch_content = (ch.content or "").strip()
        html_content = f"<h1>{ch_title}</h1>\n" + ch_content.replace("\n", "<br/>\n")
        item = epub.EpubHtml(
            title=ch_title,
            file_name=f"chap_{ch.id}.xhtml",
            lang="zh",
        )
        item.content = f"<html><body>{html_content}</body></html>"
        book.add_item(item)
        epub_items.append(item)
        spine.append(item)

    book.toc = epub_items
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    book.spine = spine

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    buf.seek(0)
    return buf


# ── DOCX ──────────────────────────────────────────────

async def _generate_docx(novel: Novel, chapters: List[Chapter]) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(12)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"《{novel.title}》")
    run.font.size = Pt(22)
    run.bold = True

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run(f"作者：{novel.author.username}")
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    for ch in chapters:
        doc.add_heading(ch.title or f"第{ch.order_index}章", level=1)
        content = (ch.content or "").strip()
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()  # separator

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── HTML ──────────────────────────────────────────────

async def _generate_html(novel: Novel, chapters: List[Chapter]) -> io.BytesIO:
    toc_items = "".join(
        f'<li><a href="#ch{ch.id}">{ch.title}</a></li>'
        for ch in chapters
    )

    body_parts = []
    for ch in chapters:
        content = (ch.content or "").strip()
        html_content = content.replace("\n", "<br>")
        body_parts.append(
            f'<section id="ch{ch.id}">'
            f'<h2>{ch.title}</h2>'
            f'<div class="content">{html_content}</div>'
            f'</section>'
        )
    sections = "\n".join(body_parts)

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>《{novel.title}》</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", "Noto Sans SC", sans-serif;
    line-height: 1.8; color: #333; background: #fff;
    max-width: 800px; margin: 0 auto; padding: 2rem 1.5rem;
  }}
  h1 {{ text-align: center; font-size: 2rem; margin-bottom: 0.5rem; }}
  .author {{ text-align: center; color: #666; margin-bottom: 2.5rem; font-size: 1.1rem; }}
  nav.toc {{ background: #f8f8f8; padding: 1.2rem 1.8rem; border-radius: 8px; margin-bottom: 2.5rem; }}
  nav.toc h3 {{ font-size: 1rem; color: #999; margin-bottom: 0.6rem; }}
  nav.toc ul {{ list-style: none; }}
  nav.toc li {{ margin: 0.3rem 0; }}
  nav.toc a {{ color: #1a73e8; text-decoration: none; }}
  section {{ margin-bottom: 2.5rem; }}
  section h2 {{ font-size: 1.4rem; margin-bottom: 1rem; border-bottom: 2px solid #eee; padding-bottom: 0.4rem; }}
  .content {{ text-indent: 2em; }}
  @media print {{
    nav.toc {{ display: none; }}
    body {{ padding: 0; max-width: none; }}
    section {{ page-break-after: always; }}
  }}
</style>
</head>
<body>
<h1>《{novel.title}》</h1>
<p class="author">作者：{novel.author.username}</p>
<nav class="toc">
  <h3>目录</h3>
  <ul>{toc_items}</ul>
</nav>
{sections}
</body>
</html>"""

    buf = io.BytesIO(html.encode("utf-8"))
    buf.seek(0)
    return buf


# ── PDF ───────────────────────────────────────────────

async def _generate_pdf(novel: Novel, chapters: List[Chapter]) -> io.BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=novel.title,
        author=novel.author.username,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CNTitle", parent=styles["Title"],
        fontName="Helvetica", fontSize=24, spaceAfter=12,
        alignment=1,
    )
    author_style = ParagraphStyle(
        "CNAuthor", parent=styles["Normal"],
        fontName="Helvetica", fontSize=14, spaceAfter=30,
        alignment=1, textColor="#666666",
    )
    heading_style = ParagraphStyle(
        "CNHeading", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=16, spaceBefore=20, spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "CNBody", parent=styles["Normal"],
        fontName="Helvetica", fontSize=11, leading=18,
        firstLineIndent=2 * cm, spaceAfter=6,
    )

    story = []
    story.append(Paragraph(f"《{novel.title}》", title_style))
    story.append(Paragraph(f"作者：{novel.author.username}", author_style))
    story.append(Spacer(1, 1 * cm))

    for ch in chapters:
        story.append(Paragraph(ch.title or f"第{ch.order_index}章", heading_style))
        content = (ch.content or "").strip()
        for para in content.split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
        story.append(PageBreak())

    doc.build(story)
    buf.seek(0)
    return buf


# ── Generator dispatch ────────────────────────────────

GENERATORS = {
    "txt": _generate_txt,
    "epub": _generate_epub,
    "docx": _generate_docx,
    "html": _generate_html,
    "pdf": _generate_pdf,
}


# ── Route ─────────────────────────────────────────────

@router.get("")
async def export_novel(
    novel_id: int,
    format: str = Query(..., description="Export format: txt, epub, docx, pdf, html"),
    chapter_ids: Optional[str] = Query(None, description="Comma-separated chapter IDs"),
    user: User = Depends(get_current_approved_user),
    db: AsyncSession = Depends(get_db),
):
    """Export a novel in the requested format.

    Returns a file download with the appropriate Content-Type and filename.
    """
    if format not in VALID_FORMATS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported format: {format}. Valid: {', '.join(sorted(VALID_FORMATS))}",
        )

    # Fetch novel with author eagerly loaded
    result = await db.execute(
        select(Novel)
        .options(selectinload(Novel.author))
        .where(Novel.id == novel_id)
    )
    novel = result.scalar_one_or_none()
    if not novel:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Novel not found")

    # Access check: owner or admin
    if novel.user_id != user.id and user.role != "admin":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not authorized")

    # Fetch chapters
    stmt = select(Chapter).where(Chapter.novel_id == novel_id)
    if chapter_ids:
        ids = [int(cid.strip()) for cid in chapter_ids.split(",") if cid.strip().isdigit()]
        if ids:
            stmt = stmt.where(Chapter.id.in_(ids))
    stmt = stmt.order_by(Chapter.order_index)
    result = await db.execute(stmt)
    chapters: List[Chapter] = list(result.scalars().all())

    generator = GENERATORS[format]
    buf = await generator(novel, chapters)

    filename = _make_filename(novel.title, format)
    return StreamingResponse(
        buf,
        media_type=CONTENT_TYPES[format],
        headers={"Content-Disposition": _disposition(filename)},
    )


# Utility: list available formats
@router.get("/formats")
async def list_formats(
    novel_id: int,
    user: User = Depends(get_current_approved_user),
    db: AsyncSession = Depends(get_db),
):
    """List supported export formats with details."""
    result = await db.execute(select(Novel).where(Novel.id == novel_id))
    novel = result.scalar_one_or_none()
    if not novel:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Novel not found")
    if novel.user_id != user.id and user.role != "admin":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Not authorized")

    return {
        "formats": [
            {"id": "txt", "label": "纯文本 TXT", "content_type": CONTENT_TYPES["txt"], "ext": "txt"},
            {"id": "epub", "label": "电子书 EPUB", "content_type": CONTENT_TYPES["epub"], "ext": "epub"},
            {"id": "docx", "label": "Word 文档 DOCX", "content_type": CONTENT_TYPES["docx"], "ext": "docx"},
            {"id": "pdf", "label": "PDF 文档", "content_type": CONTENT_TYPES["pdf"], "ext": "pdf"},
            {"id": "html", "label": "HTML 网页", "content_type": CONTENT_TYPES["html"], "ext": "html"},
        ]
    }
